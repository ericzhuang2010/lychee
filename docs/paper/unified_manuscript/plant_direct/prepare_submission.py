#!/usr/bin/env python3
"""Build a Plant Direct submission package from the unified Markdown manuscript."""

from __future__ import annotations

import hashlib
import re
import shutil
import textwrap
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUTPUT_DIR = Path(__file__).resolve().parent
MANUSCRIPT_DIR = OUTPUT_DIR.parent
PROJECT_ROOT = OUTPUT_DIR.parents[3]
SOURCE_MD = MANUSCRIPT_DIR / "manuscript.md"
SOURCE_SUPPLEMENT = MANUSCRIPT_DIR / "lychee_unified_manuscript_supplement.zip"

TITLE = (
    "Cultivar-dependent transcriptional responses of lychee to "
    "Peronophythora litchii: a registered genome-wide analysis"
)
AUTHOR = "Eric Zhuang"
AFFILIATION = "NYU Langone Health, New York, NY, USA"
EMAIL = "eric.zhuang@nyulangone.org"
ORCID = "0009-0001-9050-0214"
ZENODO_DOI = "https://doi.org/10.5281/zenodo.22240717"

FIGURES = [
    (
        1,
        "figure1_study_design_qc.pdf",
        "Figure_1_study_design_and_QC.pdf",
    ),
    (
        2,
        "figure2_discovery_legacy.pdf",
        "Figure_2_discovery_and_legacy_audit.pdf",
    ),
    (
        3,
        "figure3_robustness_external.pdf",
        "Figure_3_robustness_and_external_evaluation.pdf",
    ),
    (
        4,
        "figure4_pathways_signatures.pdf",
        "Figure_4_pathway_and_signature_evaluation.pdf",
    ),
    (
        5,
        "figure5_transcript_usage.pdf",
        "Figure_5_transcript_usage.pdf",
    ),
    (
        6,
        "figure6_orthogonal_tiers.pdf",
        "Figure_6_orthogonal_evidence_and_tiers.pdf",
    ),
]

SUPPLEMENTARY_FIGURES = [
    (
        "Figure S1",
        PROJECT_ROOT / "results/figures/FigureS1_replicate_level_counts.png",
        "Replicate-level normalized counts for the two cross-context-supported genes, "
        "twelve Tier B genes, and two legacy highlights. Points are individual deposited "
        "libraries; bars show cultivar-treatment medians.",
    ),
    (
        "Figure S2",
        PROJECT_ROOT / "results/figures/FigureS2_power_analysis.png",
        "Parametric detection probability for cultivar-by-infection effects under "
        "genome-wide discovery and candidate-family external adjustment. Curves show the "
        "overall result and mean-expression quartiles; the dashed line marks 80% power.",
    ),
    (
        "Figure S3",
        MANUSCRIPT_DIR / "figures/figureS3_exploratory_signature.png",
        "Quarantined PRJNA1090613 signed-signature estimate, retained as exploratory "
        "rather than confirmatory evidence.",
    ),
]

SUPPLEMENTAL_ITEMS = [
    "Table S1. Biological-unit registry.",
    "Table S2. Per-library quality control.",
    "Table S3. Genome-wide discovery statistics.",
    "Table S4. Robustness file manifest.",
    "Table S5. All external frozen tests.",
    "Table S6. Pathway and signature tests.",
    "Table S7. Conditional differential-transcript-usage results.",
    "Table S8. Candidate annotation and orthology evidence.",
    "Table S9. Small-RNA reference-gate outcome.",
    "Table S10. Motif-background tests and sensitivity results.",
    "Table S11. Accession-aware evidence registry.",
    "Table S12. Scripts, environments, and command inventory.",
    "Table S13. Protocol amendment and deviation log.",
    "Table S14. Legacy within-cultivar audit.",
    "Table S15. Controlled promoter-motif background comparison.",
    "Table S16a. Reconstructed dataset-search queries.",
    "Table S16b. Dataset eligibility decisions.",
    "Table S17. Exact software versions.",
    "Table S18. Power-simulation minimum detectable effects.",
    "Figure S1. Replicate-level normalized counts.",
    "Figure S2. Conditional parametric power curves.",
    "Figure S3. Exploratory signed-signature estimate.",
    "Data S1. Tab-separated source data for all main and supplementary analytical figures.",
]

INLINE_TOKEN = re.compile(
    r"(\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*)|`[^`]+`|\[[^]]+\]\([^)]+\))"
)
IMAGE_RE = re.compile(r"^!\[(?P<caption>.+)]\((?P<path>[^)]+)\)$")
HEADING_RE = re.compile(r"^(#{1,3})\s+(.+)$")
REFERENCE_RE = re.compile(r"^\d+\.\s+")


def plain_markdown(value: str) -> str:
    value = re.sub(r"\[([^]]+)]\([^)]+\)", r"\1", value)
    return value.replace("**", "").replace("*", "").replace("`", "")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, value: str) -> None:
    """Add the small Markdown subset used by the manuscript as Word runs."""
    cursor = 0
    for match in INLINE_TOKEN.finditer(value):
        if match.start() > cursor:
            paragraph.add_run(value[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif token.startswith("["):
            link = re.match(r"\[([^]]+)]\(([^)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2))
        cursor = match.end()
    if cursor < len(value):
        paragraph.add_run(value[cursor:])


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, end])


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(cell, top: int = 80, start: int = 80, bottom: int = 80, end: int = 80) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "start" if edge == "start" else "end" if edge == "end" else edge
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def set_keep_with_next(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    properties.append(OxmlElement("w:keepNext"))


def configure_document(document: Document, *, line_numbers: bool) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    if line_numbers:
        section_properties = section._sectPr
        line_numbering = OxmlElement("w:lnNumType")
        line_numbering.set(qn("w:countBy"), "1")
        line_numbering.set(qn("w:start"), "1")
        line_numbering.set(qn("w:restart"), "continuous")
        section_properties.append(line_numbering)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)

    for name, size, before, after in (
        ("Title", 16, 0, 12),
        ("Heading 1", 14, 12, 6),
        ("Heading 2", 12, 10, 4),
    ):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Figure Callout" not in [style.name for style in document.styles]:
        callout = document.styles.add_style("Figure Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.base_style = document.styles["Normal"]
        callout.font.name = "Times New Roman"
        callout.font.bold = True
        callout.font.color.rgb = RGBColor(89, 89, 89)
        callout.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        callout.paragraph_format.space_before = Pt(9)
        callout.paragraph_format.space_after = Pt(9)
        callout.paragraph_format.line_spacing = 1

    if line_numbers:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Page ")
        add_field(footer.add_run(), " PAGE ")


def add_title_page(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.15
    add_inline(
        paragraph,
        "Cultivar-dependent transcriptional responses of lychee to "
        "*Peronophythora litchii*: a registered genome-wide analysis",
    )
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)
        run.font.bold = True
    paragraph.paragraph_format.space_after = Pt(24)

    for line, bold in (
        (AUTHOR, True),
        (AFFILIATION, False),
        ("Corresponding author: Eric Zhuang", False),
        (f"Email: {EMAIL}", False),
        (f"ORCID iD: {ORCID}", False),
    ):
        item = document.add_paragraph()
        item.alignment = WD_ALIGN_PARAGRAPH.CENTER
        item.paragraph_format.line_spacing = 1.15
        run = item.add_run(line)
        run.bold = bold

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.line_spacing = 1.15
    note.add_run("Article type: Original Research Article")


def extract_abstract(lines: list[str]) -> str:
    start = lines.index("## Abstract") + 1
    for line in lines[start:]:
        if line.strip() and not line.startswith("**Keywords:"):
            return line.strip()
    raise ValueError("Abstract not found")


def extract_figure_legends(lines: list[str]) -> dict[int, str]:
    legends: dict[int, str] = {}
    for line in lines:
        match = IMAGE_RE.match(line.strip())
        if not match:
            continue
        caption = match.group("caption")
        number = re.match(r"Figure ([1-6])\.\s*(.+)", caption)
        if number:
            legends[int(number.group(1))] = caption
    if sorted(legends) != list(range(1, 7)):
        raise ValueError(f"Expected Figure 1-6 legends, found {sorted(legends)}")
    return legends


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, source_row in enumerate(rows):
        row = table.rows[row_index]
        if row_index == 0:
            set_repeat_table_header(row)
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                shade_cell(cell, "D9EAF7")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_after = Pt(0)
            if column_index < len(source_row):
                add_inline(paragraph, source_row[column_index])
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(8)
                if row_index == 0:
                    run.bold = True
    following = document.add_paragraph()
    following.paragraph_format.line_spacing = 1
    following.paragraph_format.space_after = Pt(0)


def add_main_body(document: Document, lines: list[str]) -> None:
    start = lines.index("## 1. Introduction")
    end = lines.index("## Declarations")
    index = start
    pending_caption = False

    while index < end:
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue

        heading = HEADING_RE.match(stripped)
        if heading:
            level = len(heading.group(1))
            style = "Heading 1" if level == 2 else "Heading 2"
            paragraph = document.add_paragraph(style=style)
            add_inline(paragraph, heading.group(2))
            index += 1
            continue

        image = IMAGE_RE.match(stripped)
        if image:
            number = re.match(r"Figure ([1-6])\.", image.group("caption"))
            if number:
                document.add_paragraph(
                    f"[Insert Figure {number.group(1)} near here]", style="Figure Callout"
                )
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            pending_caption = False
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < end:
            candidate = lines[index].strip()
            if not candidate:
                break
            if HEADING_RE.match(candidate) or IMAGE_RE.match(candidate) or candidate.startswith("|"):
                break
            paragraph_lines.append(candidate)
            index += 1
        value = " ".join(paragraph_lines)
        paragraph = document.add_paragraph()
        add_inline(paragraph, value)
        if value.startswith("**Table "):
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            set_keep_with_next(paragraph)
            pending_caption = True
        elif pending_caption:
            pending_caption = False


def add_declarations_and_data(document: Document) -> None:
    sections = [
        ("Acknowledgments", "The author has no acknowledgments to declare."),
        ("Funding", "This work received no external funding."),
        (
            "Author Contributions",
            "E.Z. conceived the study, performed all analyses, and wrote the manuscript. "
            "The author approved the submitted version and is accountable for the work.",
        ),
        (
            "Conflict of Interest",
            "The author declares no conflicts of interest.",
        ),
        (
            "Data Availability",
            "All analyzed sequencing data are public under PRJNA830488/GSE201243, "
            "PRJNA450886, PRJNA922966/GSE222651, PRJNA922965/GSE222650, and "
            "PRJNA1090613/GSE262200. Supplementary Tables S1-S18, Figures S1-S3, and "
            f"tab-separated figure source data are archived under CC BY 4.0 at {ZENODO_DOI}.",
        ),
        (
            "Accession Numbers",
            "Sequence data analyzed in this article can be found in the NCBI SRA and GEO "
            "repositories under PRJNA830488 (GSE201243), PRJNA450886, PRJNA922966 "
            "(GSE222651), PRJNA922965 (GSE222650), and PRJNA1090613 (GSE262200). "
            f"Supplemental data are permanently archived at {ZENODO_DOI}.",
        ),
    ]
    for heading, body in sections:
        document.add_paragraph(heading, style="Heading 1")
        paragraph = document.add_paragraph()
        add_inline(paragraph, body)

    document.add_paragraph("Supplemental Data", style="Heading 1")
    intro = document.add_paragraph()
    add_inline(
        intro,
        f"The following items are archived in the supplement at {ZENODO_DOI}:",
    )
    for item in SUPPLEMENTAL_ITEMS:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(item)


def extract_references(lines: list[str]) -> list[str]:
    start = lines.index("## References") + 1
    references: list[str] = []
    current: list[str] = []
    for line in lines[start:]:
        stripped = line.strip()
        if not stripped:
            continue
        if REFERENCE_RE.match(stripped):
            if current:
                references.append(" ".join(current))
            current = [stripped]
        elif current:
            current.append(stripped)
    if current:
        references.append(" ".join(current))
    return references


def add_references(document: Document, references: list[str]) -> None:
    heading = document.add_paragraph("References", style="Heading 1")
    heading.paragraph_format.page_break_before = True
    for reference in references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(4)
        add_inline(paragraph, reference)


def add_figure_legends(document: Document, legends: dict[int, str]) -> None:
    heading = document.add_paragraph("Figure Legends", style="Heading 1")
    heading.paragraph_format.page_break_before = True
    for number in range(1, 7):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(8)
        prefix = f"Figure {number}."
        caption = legends[number]
        if not caption.startswith(prefix):
            raise ValueError(f"Unexpected caption for Figure {number}: {caption}")
        paragraph.add_run(prefix).bold = True
        paragraph.add_run(caption[len(prefix) :])


def build_manuscript() -> Path:
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    abstract = extract_abstract(lines)
    if len(re.findall(r"\b[\w'-]+\b", abstract)) > 500:
        raise ValueError("Abstract exceeds Plant Direct's 500-word limit")
    legends = extract_figure_legends(lines)
    references = extract_references(lines)

    document = Document()
    configure_document(document, line_numbers=True)
    properties = document.core_properties
    properties.title = TITLE
    properties.author = AUTHOR
    properties.subject = "Original Research Article submitted to Plant Direct"
    properties.keywords = (
        "Litchi chinensis; Peronophythora litchii; RNA-seq; cultivar-by-infection "
        "interaction; preregistered analysis; cross-context evaluation"
    )

    add_title_page(document)
    abstract_heading = document.add_paragraph("Abstract", style="Heading 1")
    abstract_heading.paragraph_format.page_break_before = True
    abstract_paragraph = document.add_paragraph()
    add_inline(abstract_paragraph, abstract)
    keywords = document.add_paragraph()
    keywords.paragraph_format.line_spacing = 1.15
    keywords.add_run("Keywords: ").bold = True
    add_inline(
        keywords,
        "*Litchi chinensis*; *Peronophythora litchii*; RNA-seq; "
        "cultivar-by-infection interaction; preregistered analysis; cross-context evaluation",
    )

    add_main_body(document, lines)
    add_declarations_and_data(document)
    add_references(document, references)
    add_figure_legends(document, legends)

    output = OUTPUT_DIR / "Plant_Direct_manuscript.docx"
    document.save(output)
    return output


def build_cover_letter() -> Path:
    document = Document()
    configure_document(document, line_numbers=False)
    properties = document.core_properties
    properties.title = f"Cover letter: {TITLE}"
    properties.author = AUTHOR
    properties.subject = "Submission to Plant Direct"

    normal = document.styles["Normal"]
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for line in (
        date.today().strftime("%B %-d, %Y"),
        "Editors",
        "Plant Direct",
        "Wiley",
    ):
        document.add_paragraph(line)

    document.add_paragraph("Dear Editors,")

    paragraphs = [
        (
            "Please consider my manuscript, “Cultivar-dependent transcriptional responses "
            "of lychee to *Peronophythora litchii*: a registered genome-wide analysis,” for "
            "publication in Plant Direct as an Original Research Article."
        ),
        (
            "This study addresses a central challenge in plant-pathogen transcriptomics: "
            "distinguishing cultivar-dependent infection responses from effects selected within "
            "a single small cohort. It combines a genome-wide cultivar-by-infection analysis, "
            "a prospectively registered confirmatory workflow, independent statistical and "
            "quantification checks, and prespecified cross-context evaluation in public lychee "
            "RNA-seq cohorts. The analysis reports supported, null, and contradictory results "
            "under a deterministic evidence framework, including an empty highest-confidence "
            "tier where the evidence did not justify stronger claims."
        ),
        (
            "The work fits Plant Direct's broad plant-science and bioinformatics scope and its "
            "emphasis on methodologically sound research. Its main contribution is a transparent "
            "candidate resource with explicit evidence boundaries, together with a practical "
            "demonstration that plausible single-cohort candidates may not survive a formal "
            "genome-wide interaction test."
        ),
        (
            "The manuscript is not under consideration by another journal and has not been "
            "published as a peer-reviewed article. I am the sole author, have approved the "
            "submitted version, and accept responsibility for the integrity of the work. The "
            "study is a reanalysis of public sequencing data and involved no new human or animal "
            "participants. It received no external funding, and I declare no conflicts of interest."
        ),
        (
            "All underlying sequencing data are publicly accessioned. The complete supplementary "
            f"tables, supplementary figures, and figure source data are archived at {ZENODO_DOI}, "
            "with accession details provided in the manuscript."
        ),
        (
            "Thank you for considering this manuscript. I believe its emphasis on sound inference, "
            "transparent negative results, and reproducible computational plant biology will be "
            "useful to Plant Direct's readership."
        ),
    ]
    for value in paragraphs:
        paragraph = document.add_paragraph()
        add_inline(paragraph, value)

    for line in ("Sincerely,", AUTHOR, AFFILIATION, EMAIL, f"ORCID iD: {ORCID}"):
        document.add_paragraph(line)

    output = OUTPUT_DIR / "Plant_Direct_cover_letter.docx"
    document.save(output)
    return output


def fit_image(image: Image.Image, maximum_width: int, maximum_height: int) -> Image.Image:
    fitted = image.copy()
    fitted.thumbnail((maximum_width, maximum_height), Image.Resampling.LANCZOS)
    if fitted.mode != "RGB":
        background = Image.new("RGB", fitted.size, "white")
        if "A" in fitted.getbands():
            background.paste(fitted, mask=fitted.getchannel("A"))
        else:
            background.paste(fitted)
        fitted = background
    return fitted


def wrapped_lines(draw: ImageDraw.ImageDraw, text: str, font, maximum_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        candidate = " ".join(current + [word])
        if current and draw.textbbox((0, 0), candidate, font=font)[2] > maximum_width:
            lines.append(" ".join(current))
            current = [word]
        else:
            current.append(word)
    if current:
        lines.append(" ".join(current))
    return lines


def build_supplementary_figures_pdf() -> Path:
    page_width, page_height = 2550, 3300
    margin = 225
    title_font = ImageFont.truetype(
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf", 54
    )
    body_font = ImageFont.truetype(
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf", 42
    )
    pages: list[Image.Image] = []
    for label, path, caption in SUPPLEMENTARY_FIGURES:
        if not path.is_file():
            raise FileNotFoundError(path)
        page = Image.new("RGB", (page_width, page_height), "white")
        draw = ImageDraw.Draw(page)
        draw.text((margin, margin), label, font=title_font, fill="black")
        full_caption = f"{label}. {caption}"
        lines = wrapped_lines(draw, full_caption, body_font, page_width - 2 * margin)
        line_height = 52
        caption_height = line_height * len(lines)
        image_top = margin + 95
        image_height = page_height - image_top - margin - caption_height - 80
        with Image.open(path) as source:
            fitted = fit_image(source, page_width - 2 * margin, image_height)
        image_x = (page_width - fitted.width) // 2
        page.paste(fitted, (image_x, image_top))
        caption_y = image_top + fitted.height + 55
        for line in lines:
            draw.text((margin, caption_y), line, font=body_font, fill="black")
            caption_y += line_height
        pages.append(page)

    output = OUTPUT_DIR / "supporting_information/Plant_Direct_supporting_figures.pdf"
    output.parent.mkdir(parents=True, exist_ok=True)
    pages[0].save(
        output,
        "PDF",
        resolution=300.0,
        save_all=True,
        append_images=pages[1:],
        quality=95,
    )
    return output


def copy_payload_files() -> list[Path]:
    figure_dir = OUTPUT_DIR / "figures"
    support_dir = OUTPUT_DIR / "supporting_information"
    figure_dir.mkdir(parents=True, exist_ok=True)
    support_dir.mkdir(parents=True, exist_ok=True)
    outputs: list[Path] = []
    for _, source_name, target_name in FIGURES:
        source = MANUSCRIPT_DIR / "figures" / source_name
        target = figure_dir / target_name
        shutil.copy2(source, target)
        outputs.append(target)
    supplement_target = support_dir / SOURCE_SUPPLEMENT.name
    shutil.copy2(SOURCE_SUPPLEMENT, supplement_target)
    outputs.append(supplement_target)
    return outputs


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def write_manifest(paths: list[Path]) -> Path:
    output = OUTPUT_DIR / "UPLOAD_FILE_MANIFEST_SHA256.tsv"
    lines = ["sha256\tbytes\tfile"]
    for path in sorted(paths, key=lambda item: str(item.relative_to(OUTPUT_DIR))):
        lines.append(f"{sha256(path)}\t{path.stat().st_size}\t{path.relative_to(OUTPUT_DIR)}")
    output.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return output


def main() -> None:
    manuscript = build_manuscript()
    cover_letter = build_cover_letter()
    supplementary_figures = build_supplementary_figures_pdf()
    copied = copy_payload_files()
    write_manifest([manuscript, cover_letter, supplementary_figures, *copied])
    print(f"Built Plant Direct package in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
